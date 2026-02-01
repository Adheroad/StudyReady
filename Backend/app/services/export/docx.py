"""DOCX export service for CBSE question papers with precise layout matching."""

import re
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.logging import get_logger

logger = get_logger(__name__)


def _set_cell_border(cell, **kwargs):
    """
    Set cell borders for a docx table cell.
    Usage: _set_cell_border(cell, top={"sz": 12, "color": "000000", "val": "single"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))


def generate_docx(
    paper_content: str,
    subject: str,
    grade: str,
    total_marks: int,
    output_path: Optional[str] = None,
) -> bytes:
    """Generate DOCX with high-fidelity CBSE layout."""
    logger.info("Generating high-fidelity DOCX", subject=subject)
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- Header Section (Series & Set) ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(7)
    
    # Series (Left)
    series_match = re.search(r"Series:\s*([^\n]+)", paper_content)
    series = series_match.group(1).strip() if series_match else "WXY4Z"
    cell_l = header_table.cell(0, 0)
    p_l = cell_l.paragraphs[0]
    p_l.add_run(f"Series : {series}").bold = True
    _set_cell_border(cell_l, top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"}, 
                     start={"sz": 4, "val": "single"}, end={"sz": 4, "val": "single"})

    # Set (Right)
    set_match = re.search(r"Set:\s*([^\n]+)", paper_content)
    set_num = set_match.group(1).strip() if set_match else "4"
    cell_r = header_table.cell(0, 1)
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_r.add_run(f"SET-{set_num}")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '000000')
    cell_r._tc.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph() # Spacer

    # --- QP Code ---
    code_match = re.search(r"Code:\s*([^\n]+)", paper_content)
    code = code_match.group(1).strip() if code_match else "72/708"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("प्रश्न-पत्र कोड / Q.P. Code ").font.size = Pt(10)
    p.add_run(code).font.size = Pt(20)
    p.runs[-1].bold = True

    # --- Roll No ---
    roll_table = doc.add_table(rows=1, cols=11)
    roll_table.style = 'Table Grid'
    cell_label = roll_table.cell(0, 0)
    cell_label.text = "रोल नं.\nRoll No."
    cell_label.paragraphs[0].runs[0].font.size = Pt(9)
    # Clear borders for label cell
    _set_cell_border(cell_label, top={"val": "nil"}, start={"val": "nil"}, bottom={"val": "nil"}, end={"val": "nil"})

    # --- Subject Title ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{subject}\n{subject} (Theory)")
    run.bold = True
    run.font.size = Pt(14)

    # --- Meta ---
    p = doc.add_paragraph()
    run = p.add_run(f"निर्धारित समय: 2 घण्टे / Time allowed: 2 hours")
    run.font.size = Pt(10)
    run.bold = True
    run.add_text(f"\t\t\tअधिकतम अंक: {total_marks} / Maximum Marks: {total_marks}")

    # --- Questions ---
    lines = paper_content.split('\n')
    for line in lines:
        if line.startswith('## SECTION'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        elif line.startswith('SECTION') or line.startswith('खण्ड'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
        elif line.startswith('•'):
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
        elif line[0:1].isdigit() and '.' in line[:4]:
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
        else:
            doc.add_paragraph(line)

    # Save
    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(docx_bytes)
            
    return docx_bytes
